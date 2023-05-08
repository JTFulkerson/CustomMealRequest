# ----------------------------------------------------------------------------
# Created By: John Fulkerson
# With help from: Brennan Gallamoza
# Created Date: 12/6/2022
# ----------------------------------------------------------------------------

from ctypes.wintypes import SERVICE_STATUS_HANDLE
from flask import Flask, redirect, render_template, request, url_for
from dataclasses import dataclass
import os
from stat import S_IWUSR, S_IRUSR, S_IRGRP, S_IROTH
import shutil
from docx2pdf import convert
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt

# dates
from datetime import datetime, timedelta

# email
import smtplib
from pathlib import Path
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email.utils import formatdate, COMMASPACE
from email import encoders


@dataclass
class Person:
    """
    Creates a person with name(string), phone(string), email(), restriction(string)
    """
    name: str
    phone: str
    email: str
    restriction: str


@dataclass
class Meal:
    """
    Creates a meal with fields, type(string), time(string), location(string),
    food(list)
    """
    type: str
    time: str
    location: str
    food: list


@dataclass
class Request:
    """Contains all 3 meals, Breakfast, Lunch, and Dinner
    """
    breakfast: Meal
    lunch: Meal
    dinner: Meal
    recipients: list[str]


def list_food(food: str) -> list[str]:
    """Takes a string of food items and returns a list of food items.

    Args:
        food (str): List of food items separated by a period.

    Returns:
        list[str]: List of food items.
    """
    lst = []
    # checks to see if period is at end of string, if so removes it
    if food != "":
        if food[-1] == ".":
            food = food[:-1]
        lst = food.split(".")  # splits string into list
        for i in range(len(lst)):
            lst[i] = ' - ' + lst[i].strip()  # adds dash and removes whitespace
    return lst


def layout_order(meal: Meal) -> str:
    """Takes a meal and outputs the food items in the string with
       line breaks between items.

    Args:
        meal (Meal): _description_

    Returns:
        str: _description_
    """
    order_string = ""
    for item in meal.food:
        order_string += item + "\n"
    return order_string


def convert_location_to_email(loc: str) -> str:
    if loc == "Russell":
        return "russelldininghall@udel.edu"
    elif loc == "Caesar Rodney":
        return "rodneydiningffco@udel.edu"
    elif loc == "Pencader":
        return "pencaderdininghall@udel.edu"
    else:
        return ""


def create_recipient_list(person: Person, the_order: Request) -> list[str]:
    """Creates list of email addresses corresponding to dinning halls that
       need the meal request.

    Args:
        the_order (Request): Meal that was just created by the user.

    Returns:
        list[str]: Email addresses corresponding to dinning halls that need
                   the meal request
    """
    recipient_list = [person.email]
    if person.name.lower() != 'test':
        if convert_location_to_email(the_order.breakfast.location) != "":
            recipient_list.append(convert_location_to_email(
                the_order.breakfast.location))
        if convert_location_to_email(the_order.lunch.location) != "":
            recipient_list.append(
                convert_location_to_email(the_order.lunch.location))
        if convert_location_to_email(the_order.dinner.location) != "":
            recipient_list.append(
                convert_location_to_email(the_order.dinner.location))
        if not recipient_list:
            print("No meals were ordered")
            exit()
    return list(set(recipient_list))


def make_new_form(the_person: Person, the_order: Request, d: str, mdy: str, word_docx_destination: str):
    """Creates a new form from the template and fills in the information

    Args:
        the_person (Person): Person who is ordering the meal
        the_order (Request): Meal that was just created by the user.
        d (str): Day of the week
        mdy (str): Date in the format of Month Day, Year
        word_docx_destination (str): Path to store the new form

    Returns:
        None

    """
    # Gets path to form template
    form_template = "./Custom Meal Request Form.docx"
    shutil.copyfile(form_template, word_docx_destination)
    os.chmod(word_docx_destination, S_IRUSR | S_IWUSR | S_IRGRP | S_IROTH)
    document = Document(word_docx_destination)
    os.chmod(word_docx_destination, 0o777)
    table = document.tables[0]
    table.cell(0, 0).text = "Name: " + the_person.name
    table.cell(1, 0).text = "Day of Week: " + d
    table.cell(1, 3).text = "Date: " + mdy
    table.cell(2, 0).text = "Diet Restriction: " + the_person.restriction
    table.cell(0, 3).text = "Phone: " + the_person.phone

    table.cell(3, 1).text = "Time: " + \
        removeLeadingZeros(twenty_four_hour_to_twelve_hour(
            the_order.breakfast.time))
    table.cell(3, 2).text = "Location: " + the_order.breakfast.location
    table.cell(4, 0).text = layout_order(the_order.breakfast)

    table.cell(5, 1).text = "Time: " + \
        removeLeadingZeros(
            twenty_four_hour_to_twelve_hour(the_order.lunch.time))
    table.cell(5, 2).text = "Location: " + the_order.lunch.location
    table.cell(6, 0).text = layout_order(the_order.lunch)

    table.cell(7, 1).text = "Time: " + \
        removeLeadingZeros(
            twenty_four_hour_to_twelve_hour(the_order.dinner.time))
    table.cell(7, 2).text = "Location: " + the_order.dinner.location
    table.cell(8, 0).text = layout_order(the_order.dinner)

    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(16)

    document.save(word_docx_destination)


def send_email(send_from: str, name: str, send_to, the_subject: str,
               the_message, files=[], server="localhost", port=587,
               username='', password='', use_tls=True):
    """Compose and send email with provided info and attachments.

    Args:
        send_from (str): _description_
        name (str): _description_
        send_to (_type_): _description_
        the_subject (str): _description_
        the_message (_type_): _description_
        files (list, optional): _description_. Defaults to [].
        server (str, optional): _description_. Defaults to "localhost".
        port (int, optional): _description_. Defaults to 587.
        username (str, optional): _description_. Defaults to ''.
        password (str, optional): _description_. Defaults to ''.
        use_tls (bool, optional): _description_. Defaults to True.
    """
    msg = MIMEMultipart()
    msg['From'] = name + " <" + send_from + ">"
    msg['To'] = COMMASPACE.join(send_to)
    msg['Date'] = formatdate(localtime=True)
    msg['Subject'] = the_subject

    msg.attach(MIMEText(the_message))

    for path in files:
        part = MIMEBase('application', "octet-stream")
        with open(path, 'rb') as file:
            part.set_payload(file.read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition',
                        'attachment; filename={}'.format(Path(path).name))
        msg.attach(part)

    smtp = smtplib.SMTP(server, port)
    if use_tls:
        smtp.starttls()
    smtp.login(username, password)
    smtp.sendmail(send_from, send_to, msg.as_string())
    smtp.quit()


def create_calendar_event(meal: Meal, the_person: Person, d: str, mdy: str):
    """Creates a calendar event for the meal

    Args:
        meal (Meal): Meal to be created
        the_person (Person): Person who is ordering the meal
        d (str): Day of the week
        mdy (str): Date in the format of Month Day, Year
    """
    if meal.time == "":
        return
    event = {
        'summary': the_person.name + " - " + meal.location,
        'location': meal.location,
        'description': 'Meal Order',
        'start': {
            'dateTime': mdy + " " + meal.time + ":00",
            'timeZone': 'America/New_York',
        },
        'end': {
            'dateTime': mdy + " " + meal.time + ":00",
            'timeZone': 'America/New_York',
        },
        'recurrence': [
            'RRULE:FREQ=DAILY;COUNT=1'
        ],
        'attendees': [
            {'email': the_person.email},
        ],
        'reminders': {
            'useDefault': False,
            'overrides': [
                {'method': 'email', 'minutes': 24 * 60},
                {'method': 'popup', 'minutes': 10},
            ],
        },
    }

    event = SERVICE_STATUS_HANDLE.events().insert(
        calendarId='primary', body=event).execute()
    print('Event created: %s' % (event.get('htmlLink')))


def twenty_four_hour_to_twelve_hour(time: str) -> str:
    """Converts a time in 24 hour format to 12 hour format

    Args:
        time (str): Time in 24 hour format

    Returns:
        str: Time in 12 hour format
    """
    if time == "":
        return ""
    time_obj = datetime.strptime(time, "%H:%M")
    return time_obj.strftime("%I:%M %p")


def removeLeadingZeros(inputString):
    for k in range(len(inputString)):
        if inputString[k] != '0':
            outputString = inputString[k::]
            return outputString
    return ""


app = Flask(__name__, static_folder='static')


@app.route('/', methods=['GET', 'POST'])
def order_form():
    # Personal Info
    NAME = os.getenv("NAME")
    PHONE_NUMBER = os.getenv("PHONE_NUMBER")
    DIETARY_RESTRICTIONS = os.getenv("DIETARY_RESTRICTIONS")
    SCHOOL_EMAIL = os.getenv("SCHOOL_EMAIL")
    PERSONAL_EMAIL = os.getenv("PERSONAL_EMAIL")
    PERSONAL_EMAIL_PASSWORD = os.getenv("PERSONAL_EMAIL_PASSWORD")
    SERVER_NAME = os.getenv("SERVER_NAME")
    SERVER_PORT = os.getenv("SERVER_PORT")
    if request.method == 'POST':
        env_path = os.path.join('.', '.env')
        load_dotenv(env_path)

        date_str = request.form.get("date")
        date_obj = datetime.strptime(date_str, "%Y-%m-%d")
        d = date_obj.strftime("%A")
        mdy = date_obj.strftime("%m/%d/%Y")
        m_d_y = date_obj.strftime("%m-%d-%Y")  # prints theDate in m_d_y

        # get the form data
        person = Person(NAME, PHONE_NUMBER, SCHOOL_EMAIL, DIETARY_RESTRICTIONS)

        # gets path to the previous_meal_requests folder and names file
        word_docx_destination = "./previous_meal_requests/" + \
            m_d_y + " " + person.name + ".docx"
        # gets path to the previous_meal_requests folder and names file
        pdf_destination = "./previous_meal_requests/" + \
            m_d_y + " " + person.name + ".pdf"

        # collect the order
        order = order = Request(Meal("breakfast", request.form.get('breakfast_time'), request.form.get('breakfast_location'), list_food(request.form.get('breakfast_food'))),
                                Meal("lunch", request.form.get('lunch_time'), request.form.get(
                                    'lunch_location'), list_food(request.form.get('lunch_food'))),
                                Meal("dinner", request.form.get('dinner_time'), request.form.get('dinner_location'), list_food(request.form.get('dinner_food'))), [])

        # make the form
        make_new_form(person, order, d, mdy, word_docx_destination)

        recipients = create_recipient_list(person, order)
        convert(word_docx_destination)
        os.remove(word_docx_destination)

        # send the email
        if person.name.lower() == 'test':
            subject = "TEST CUSTOM MEAL REQUEST - " + person.name + " - " + mdy
        else:
            subject = "CUSTOM MEAL REQUEST - " + person.name + " - " + mdy
        body = "Good morning, here is my custom meal request for today. Thank you!"
        send_email(PERSONAL_EMAIL, NAME, recipients, subject, body,
                   [pdf_destination], SERVER_NAME, SERVER_PORT,
                   PERSONAL_EMAIL, PERSONAL_EMAIL_PASSWORD)

        return render_template('confirmation', name=NAME, date=date_str, breakfast_time=order.breakfast.time, breakfast_location=order.breakfast.location, breakfast_food=order.breakfast.food, lunch_time=order.lunch.time, lunch_location=order.lunch.location, lunch_food=order.lunch.food, dinner_time=order.dinner.time, dinner_location=order.dinner.location, dinner_food=order.dinner.food)
    else:
        return render_template('form.html', name=NAME)


@app.route('/success')
def success():
    return render_template('success.html')


if __name__ == '__main__':
    port_num = 3000
    print("Starting Flask Server")
    try:
        app.run(host="localhost", port=port_num, debug=True)
    except:
        for i in range(10):
            try:
                port_num += 1
                app.run(host="localhost", port=port_num, debug=True)
                break
            except:
                continue
    finally:
        print("Flask Server Started")
