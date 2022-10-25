# ----------------------------------------------------------------------------
# Created By: John Fulkerson
# With help from: Brennan Gallamoza
# Created Date: 10/4/2022
# ----------------------------------------------------------------------------

from dataclasses import dataclass
import os
from stat import S_IWUSR, S_IRUSR, S_IRGRP, S_IROTH
import shutil
from xml.dom.minidom import Document
from docx2pdf import convert
from dotenv import load_dotenv

# dates
from datetime import datetime, timedelta

# email
import smtplib
from pathlib import Path
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email.utils import formatdate, COMMASPACE
from email import encoders

"""
Need to install:
pip install python-docx
pip install docx2pdf
pip install python-dotenv
"""


@dataclass
class Person:
    """
    Creates a person with name(string), phone(string), restriction(string)
    """
    name: str
    phone: str
    restriction: str


@dataclass
class Meal:
    """
    Creates a meal with fields, type(string), time(string), location(string), food(list)
    """
    type: str
    time: str
    location: str
    food: list


@dataclass
class Request:
    """Contains all 3 meals, Breakfast, Lunch, and Dinner
    """
    breakfast: Meal
    lunch: Meal
    dinner: Meal
    recipients: list[str]


def want_meal(meal: str) -> bool:
    """Intended to return True or False whether the user wants certain meal or not.

    Returns:
        bool: _description_
    """
    user = input("Do you want to special order " + meal + "? \n").lower()
    if user == "yes":
        return True
    elif user == "no":
        return False
    else:
        print("Please only type yes or no")
        return want_meal(meal)


def what_dining_hall(meal: str) -> str:
    """Intended to enter meal and then user is prompted what dining hall they
       would like that meal in. The dining hall they choose is then returned.

    Args:
        meal (str): _description_

    Returns:
        str: _description_
    """
    while True:
        d1a = input("Which dining hall would you like " + meal +
                    " in: \n A) Russell B) Caesar Rodney C) Pencader \n").lower()
        if d1a == "a":
            return "Russell"
        elif d1a == "b":
            return "Caesar Rodney"
        elif d1a == "c":
            return "Pencader"
        else:
            print("Please enter either A, B, or C")
            return what_dining_hall(meal)


def collect_order(breakfast: bool, lunch: bool, dinner: bool) -> Request:
    """Takes order, time, location and writes into a Request type which is returned.

    Args:
        breakfast (bool): _description_
        lunch (bool): _description_
        dinner (bool): _description_

    Returns:
        Request: _description_
    """
    new_order = Request(Meal("breakfast", "", "", []), Meal("lunch", "", "", []), Meal("dinner", "", "", []), [])
    if breakfast:
        new_order.breakfast.food = collect_food("breakfast")
        new_order.breakfast.time = input("What time would you like your breakfast ready? ")
        new_order.breakfast.location = what_dining_hall("breakfast")

    if lunch:
        new_order.lunch.food = collect_food("lunch")
        new_order.lunch.time = input("What time would you like your lunch ready? ")
        new_order.lunch.location = what_dining_hall("lunch")

    if dinner:
        new_order.dinner.food = collect_food("dinner")
        new_order.dinner.time = input("What time would you like your dinner ready? ")
        new_order.dinner.location = what_dining_hall("dinner")
    return new_order


def collect_food(meal_type: str) -> list[str]:
    """Prompts user with how many items they would like to order then turn all the items into a list.

    Args:
        meal_type (str): _description_

    Returns:
        list[str]: _description_
    """
    lst = []

    # number of elements as input
    number_food_items = input("Enter the number of items you would like to order for " + meal_type + "?: ")
    if number_food_items.isdigit():
        n = int(number_food_items)
    else:
        print("Please enter a whole number")
        return collect_food(meal_type)
    # iterating till the range
    for i in range(0, n):
        ele = input("Type one item: ")

        lst.append(ele)  # adding the element

    return lst


def recognize(name: str, number: str, restriction: str) -> Person:
    """Checks to see if Person is stored locally

    Args:
        name (str): _description_
        number (str): _description_
        restriction (str): _description_

    Returns:
        Person: _description_
    """
    new_person = Person(name, number, restriction)
    temp = input("Are you " + new_person.name + "? ").lower()
    if temp == "no":
        new_person.name = input("What is your first and last name? ")
        new_person.phone = input("What is your phone number? ")
        new_person.restriction = input("What are your dietary restrictions? ")
        return new_person
    elif temp == "yes":
        return new_person
    else:
        print("Please only type yes or no ")
        recognize(name, number, restriction)


def layout_order(meal: Meal) -> str:
    """Takes a meal and outputs the food items in the string with line breaks between items.

    Args:
        meal (Meal): _description_

    Returns:
        str: _description_
    """
    order_string = ""
    for item in meal.food:
        order_string += item + "\n"
    return order_string


def convert_pdf(path: str):
    """Tells user to convert file to pdf and waits for it to happen.

    Args:
        path (str): Path to docx that needs to be converted
    """
    convert(path)
    os.remove(path)


def convert_location_to_email(loc: str) -> str:
    if loc == "Russell":
        return "russelldininghall@udel.edu"
    elif loc == "Caesar Rodney":
        return "rodneydiningffco@udel.edu"
    elif loc == "Pencader":
        return "pencaderdininghall@udel.edu"
    else:
        return ""


def create_recipient_list(the_order: Request) -> list[str]:
    """Creates list of email addresses corresponding to dinning halls that need the meal request.

    Args:
        the_order (Request): _description_

    Returns:
        list[str]: _description_
    """
    recipient_list = [SCHOOL_EMAIL]
    if convert_location_to_email(the_order.breakfast.location) != "":
        recipient_list.append(convert_location_to_email(the_order.breakfast.location))
    if convert_location_to_email(the_order.lunch.location) != "":
        recipient_list.append(convert_location_to_email(the_order.lunch.location))
    if convert_location_to_email(the_order.dinner.location) != "":
        recipient_list.append(convert_location_to_email(the_order.dinner.location))
    if not recipient_list:
        print("No meals were ordered")
        exit()
    return list(set(recipient_list))


def make_new_form(the_person: Person, the_order: Request):
    form_template = "./Custom Meal Request Form.docx"  # Gets path to form template
    shutil.copyfile(form_template, word_docx_destination)
    os.chmod(word_docx_destination, S_IRUSR | S_IWUSR | S_IRGRP | S_IROTH)
    document = Document(word_docx_destination)
    table = document.tables[0]
    table.cell(0, 0).text = "Name: " + the_person.name
    table.cell(1, 0).text = "Day of Week: " + d
    table.cell(1, 3).text = "Date: " + mdy
    table.cell(2, 0).text = "Diet Restriction: " + the_person.restriction
    table.cell(0, 3).text = "Phone: " + the_person.phone

    table.cell(3, 1).text = "Time: " + the_order.breakfast.time
    table.cell(3, 2).text = "Location: " + the_order.breakfast.location
    table.cell(4, 0).text = layout_order(the_order.breakfast)

    table.cell(5, 1).text = "Time: " + the_order.lunch.time
    table.cell(5, 2).text = "Location: " + the_order.lunch.location
    table.cell(6, 0).text = layout_order(the_order.lunch)

    table.cell(7, 1).text = "Time: " + the_order.dinner.time
    table.cell(7, 2).text = "Location: " + the_order.dinner.location
    table.cell(8, 0).text = layout_order(the_order.dinner)

    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(16)

    document.save(word_docx_destination)


def send_email(send_from, name, send_to, the_subject, the_message, files=[],
               server="localhost", port=587, username='', password='', use_tls=True):
    """Compose and send email with provided info and attachments.

    Args:
        send_from (_type_): _description_
        name (_type_): _description_
        send_to (_type_): _description_
        the_subject (_type_): _description_
        the_message (_type_): _description_
        files (list, optional): _description_. Defaults to [].
        server (str, optional): _description_. Defaults to "localhost".
        port (int, optional): _description_. Defaults to 587.
        username (str, optional): _description_. Defaults to ''.
        password (str, optional): _description_. Defaults to ''.
        use_tls (bool, optional): _description_. Defaults to True.
    """
    msg = MIMEMultipart()
    msg['From'] = name + " <" + send_from + ">"
    msg['To'] = COMMASPACE.join(send_to)
    msg['Date'] = formatdate(localtime=True)
    msg['Subject'] = the_subject

    msg.attach(MIMEText(the_message))

    for path in files:
        part = MIMEBase('application', "octet-stream")
        with open(path, 'rb') as file:
            part.set_payload(file.read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition',
                        'attachment; filename={}'.format(Path(path).name))
        msg.attach(part)

    smtp = smtplib.SMTP(server, port)
    if use_tls:
        smtp.starttls()
    smtp.login(username, password)
    smtp.sendmail(send_from, send_to, msg.as_string())
    smtp.quit()


# Dates
theDate = datetime.today() + timedelta(1)  # tomorrows date variable
mdy = theDate.strftime("%m/%d/%Y")  # prints theDate in m/d/y
m_d_y = theDate.strftime("%m-%d-%Y")  # prints theDate in m_d_y
d = theDate.strftime("%A")  # prints theDate in word form

if __name__ == "__main__":
    env_path = os.path.join('.', '.env')
    load_dotenv(env_path)

    # Personal Info
    NAME = os.getenv("NAME")
    PHONE_NUMBER = os.getenv("PHONE_NUMBER")
    DIETARY_RESTRICTIONS = os.getenv("DIETARY_RESTRICTIONS")
    SCHOOL_EMAIL = os.getenv("SCHOOL_EMAIL")
    PERSONAL_EMAIL = os.getenv("PERSONAL_EMAIL")
    PERSONAL_EMAIL_PASSWORD = os.getenv("PERSONAL_EMAIL_PASSWORD")
    SERVER_NAME = os.getenv("SERVER_NAME")
    SERVER_PORT = os.getenv("SERVER_PORT")

    print("Taking order for " + d)
    person = recognize(NAME, PHONE_NUMBER, DIETARY_RESTRICTIONS)
    # gets path to the previousMealRequests folder and names file
    word_docx_destination = "./previousMealRequests/" + m_d_y + " " + person.name + ".docx"
    # gets path to the previousMealRequests folder and names file
    pdf_destination = "./previousMealRequests/" + m_d_y + " " + person.name + ".pdf"
    if os.path.exists(pdf_destination):
        print("It appears you have already ordered for today!")
        exit()
    order = collect_order(want_meal('breakfast'), want_meal('lunch'), want_meal('dinner'))

    # file creation and editing
    from docx import Document
    from docx.shared import Pt

    make_new_form(person, order)

    recipients = create_recipient_list(order)
    convert_pdf(word_docx_destination)

    # Sending Email
    subject = "CUSTOM MEAL REQUEST - " + person.name + " - " + mdy
    body = input("Type the email body: ")
    send_email(PERSONAL_EMAIL, NAME, recipients, subject, body, [pdf_destination], SERVER_NAME, SERVER_PORT,
               PERSONAL_EMAIL, PERSONAL_EMAIL_PASSWORD)

    print("Form was saved at " + pdf_destination)
